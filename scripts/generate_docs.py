"""
Generate RuralCare AI project documentation as a Word (.docx) file.
Run: python scripts/generate_docs.py
Output: RuralCare_AI_Documentation.docx (in project root)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── helpers ──────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    else:
        run.font.color.rgb = RGBColor(0x47, 0x47, 0x47)
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1A5376")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            cells[c_idx].text = str(cell_text)
            if r_idx % 2 == 1:
                tc = cells[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "EBF5FB")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)
    return table

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1C)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F7")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    return p

def add_callout(doc, text, bg="FEF9E7"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), bg)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E86AB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ── build document ────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# =============================================================================
# COVER PAGE
# =============================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("RuralCare AI")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run("Multi-Agent Rural Healthcare Assistant")
r2.font.size = Pt(20)
r2.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tagline.add_run(
    "A first-mile digital health assistant for rural and underserved communities.\n"
    "Symptom Triage  |  RAG-Grounded Guidance  |  Facility Discovery  |  Emergency Escalation"
)
r3.font.size = Pt(13)
r3.italic = True
r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = meta.add_run(
    "Version: 1.1  |  Status: MVP Complete + Cloud Deployment Ready\n"
    "Stack: Python 3.11+ · FastAPI · Streamlit · LangGraph · ChromaDB · Ollama / Claude\n"
    "Deployed: GitHub · Render · Hugging Face Spaces"
)
r4.font.size = Pt(10)
r4.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()


# =============================================================================
# SECTION P — PRESENTATION MATERIAL
# =============================================================================
add_heading(doc, "PRESENTATION MATERIAL", 1)
add_para(doc,
    "This section provides a structured write-up suitable for creating slides or a review "
    "presentation. It covers the problem statement, solution overview, tech stack, agent "
    "design, and architecture.",
    italic=True
)
add_divider(doc)
doc.add_paragraph()

# ── P.1 Problem Statement ────────────────────────────────────────────────────
add_heading(doc, "P.1  Problem Statement", 2)
add_para(doc,
    "Rural and underserved communities in India face a compounding healthcare access crisis "
    "that cannot be solved by building more hospitals alone. The core barriers are:",
    bold=False
)
doc.add_paragraph()

problem_rows = [
    ("Geographic Isolation",
     "Over 600 million rural Indians live hours away from the nearest Primary Health Centre (PHC). "
     "A trip to a doctor costs a full day of lost wages — so patients delay care until it is critical."),
    ("Chronic Workforce Shortage",
     "India has 0.7 doctors per 1,000 population in rural areas vs. a WHO minimum of 1.0. "
     "ASHA workers are the primary point of contact for millions, but carry overwhelming caseloads "
     "with limited clinical decision support."),
    ("Language and Literacy Barriers",
     "Health information is predominantly available in English or formal Hindi. "
     "Patients who speak Tamil, Bengali, Telugu, or Kannada have no access to trusted "
     "guidance in their own language."),
    ("Inability to Recognise Emergencies",
     "Patients and caregivers cannot reliably distinguish a life-threatening emergency "
     "(stroke, cardiac arrest, eclampsia) from a serious but non-critical condition. "
     "This leads to preventable deaths from delayed escalation."),
    ("Fragmented Facility Information",
     "There is no unified, up-to-date, accessible directory of government health facilities "
     "for rural patients. Patients do not know where to go, what services are available, "
     "or whether a facility can handle their condition."),
]
add_table(doc, ["Problem", "Real-World Impact"], problem_rows, col_widths=[2.0, 4.6])
doc.add_paragraph()

add_callout(doc,
    "The result: Millions of rural patients receive care too late, at the wrong facility, "
    "or not at all — not because treatment doesn't exist, but because the information gap "
    "between patient and health system is never bridged.",
    bg="FDECEA"
)
doc.add_paragraph()

# ── P.2 Solution Overview ────────────────────────────────────────────────────
add_heading(doc, "P.2  Solution Overview — What is RuralCare AI?", 2)
add_para(doc,
    "RuralCare AI is a multilingual, AI-powered first-mile health assistant that bridges the "
    "information gap between a rural patient and the formal health system. It is available "
    "24/7, works in 6 Indian languages, and requires only a basic smartphone with internet access."
)
doc.add_paragraph()
add_para(doc,
    "When a patient or caregiver describes their symptoms, RuralCare AI does five things instantly:",
    bold=True
)

solution_steps = [
    ("1. Understands", "Extracts structured clinical information from free-form text in any Indian language"),
    ("2. Triages", "Classifies urgency into Emergency / Urgent / Moderate / Mild using WHO ICD-11 principles"),
    ("3. Guides", "Delivers evidence-based health information retrieved from WHO and NHM documents — not LLM hallucinations"),
    ("4. Locates", "Finds the nearest appropriate government health facility from a database of 87 Tamil Nadu hospitals"),
    ("5. Plans", "Generates a personalised follow-up plan and a structured briefing note for the ASHA / ANM health worker"),
]
add_table(doc, ["Action", "What It Does"], solution_steps, col_widths=[1.0, 5.6])
doc.add_paragraph()

add_callout(doc,
    "Critical distinction: RuralCare AI is NOT a diagnostic tool. It does not diagnose, "
    "prescribe, or replace a doctor. It is a triage and navigation assistant — telling patients "
    "HOW URGENTLY to seek care and WHERE to go.",
    bg="FEF9E7"
)
doc.add_paragraph()

add_heading(doc, "What Makes It Different", 3)
diff_rows = [
    ("Safety-first architecture", "5 independent safety layers prevent diagnosis or prescription at code level — not just prompt level"),
    ("RAG-grounded only", "All health guidance sourced from verified WHO/NHM documents — LLM cannot improvise medical facts"),
    ("Sub-100ms emergency detection", "Life-threatening keywords detected synchronously BEFORE any LLM call — no latency risk on emergencies"),
    ("Privacy by design", "Patient names never enter the LLM — anonymised tokens used throughout; audit logs store hashes only"),
    ("LLM-agnostic", "Runs on free local Ollama (no API cost) or switches to Claude/GPT-4o with a single env variable change"),
    ("Cloud-ready", "Deployed on GitHub; render.yaml and HF Spaces config included — Render Starter plan sufficient for evaluation"),
]
add_table(doc, ["Differentiator", "How it Works"], diff_rows, col_widths=[2.2, 4.4])
doc.add_page_break()

# ── P.3 Key Tech Stack ───────────────────────────────────────────────────────
add_heading(doc, "P.3  Key Technology Stack", 2)

add_para(doc,
    "Every technology choice was made deliberately — balancing capability, cost, offline viability, "
    "and safety requirements for a rural healthcare context.",
    italic=True
)
doc.add_paragraph()

stack_rows = [
    ("Agent Orchestration", "LangGraph",
     "Explicit state machine with 8 nodes, conditional emergency routing, and a fail-safe "
     "wrapper that catches agent exceptions without crashing the pipeline."),
    ("LLM — Default (local)", "Ollama llama3.2",
     "Runs fully offline on any laptop or server. No API key, no cost. "
     "Ideal for rural/low-connectivity deployment."),
    ("LLM — Cloud", "Anthropic Claude (claude-sonnet-4-6)",
     "Activated by setting LLM_PROVIDER=claude. Used in Render and HF Spaces deployments. "
     "Same code, one environment variable change."),
    ("RAG Framework", "LangChain + ChromaDB",
     "6 vector collections (WHO, NHM, emergency protocols, symptom mapping, drug info, "
     "health schemes). Cosine similarity search. Custom chunker avoids spacy dependency."),
    ("Embeddings — Local", "Ollama nomic-embed-text (~768-dim)",
     "High-quality semantic embeddings, runs locally, no API cost."),
    ("Embeddings — Cloud", "sentence-transformers/all-MiniLM-L6-v2",
     "Used on Render and HF Spaces where Ollama is unavailable. "
     "Activated by setting EMBED_PROVIDER=huggingface."),
    ("Frontend", "Streamlit",
     "Demo UI — multilingual, text input, document upload. Runs at port 8501. "
     "Works in Kaggle notebooks and HF Spaces without modification."),
    ("Backend API", "FastAPI + uvicorn",
     "REST API with auto-generated Swagger docs. Same pipeline callable via HTTP. "
     "Runs at port 8000."),
    ("Database", "SQLite (demo) / PostgreSQL (production)",
     "Stores sessions, audit logs, facility cache (87 TN hospitals), follow-up reminders."),
    ("Translation", "langdetect + Google Translate",
     "Auto-detects language; translates to English for LLM; back-translates response "
     "to patient's language. Supports 6 Indian languages."),
    ("Deployment", "Docker + Render + Hugging Face Spaces",
     "render.yaml configures Render deployment. README.md has HF Spaces front matter. "
     "Auto-seeds data on first boot — no manual setup needed on cloud."),
    ("Version Control", "GitHub",
     "Full source at github.com/jbobbypaul-cmyk/PGD-AIML-Examples-. "
     ".env gitignored; secrets managed via Render dashboard."),
]
add_table(doc, ["Component", "Technology", "Why This Choice"], stack_rows,
          col_widths=[1.7, 1.9, 3.0])
doc.add_page_break()

# ── P.4 Agent Details ────────────────────────────────────────────────────────
add_heading(doc, "P.4  Agent Details — The 8-Agent Pipeline", 2)
add_para(doc,
    "RuralCare AI uses a LangGraph state machine where each agent is a specialist node. "
    "Agents run in sequence, sharing a single PatientState object. A fail-safe _node wrapper "
    "around every agent ensures one failure never crashes the full pipeline.",
    italic=True
)
doc.add_paragraph()

agent_details = [
    ("1", "Symptom Intake Agent",
     "First responder",
     "Extracts structured clinical data from free-form patient text. "
     "Identifies chief complaint, symptom list, duration, and severity. "
     "Also runs the emergency keyword check BEFORE invoking the LLM — "
     "if a red-flag keyword is detected (e.g. 'chest pain', 'unconscious'), "
     "the pipeline immediately routes to Emergency Escalation.",
     "LLM required"),

    ("2", "Medical Triage Agent",
     "Urgency classifier",
     "Classifies the patient's condition into one of four urgency levels: "
     "EMERGENCY (life-threatening, act now), URGENT (serious, act within 4 hours), "
     "MODERATE (needs care within 48 hours), MILD (monitor at home). "
     "Uses WHO ICD-11 severity criteria embedded in the system prompt. "
     "Conservative default: invalid or ambiguous output is treated as URGENT.",
     "LLM required"),

    ("3", "Medical RAG Agent",
     "Knowledge retriever",
     "Queries all 6 ChromaDB collections simultaneously, deduplicates results, "
     "and passes the top chunks to the LLM as context. "
     "The LLM is strictly forbidden from adding medical information beyond what "
     "was retrieved. If no documents are found (grounding_confidence=none), "
     "a safe fallback message is returned and the LLM is not called.",
     "LLM + ChromaDB"),

    ("4", "Appointment & Facility Agent",
     "Facility locator",
     "3-level hybrid lookup: (1) SQLite with 87 Tamil Nadu government hospitals — "
     "covers all 38 TN districts with case-insensitive matching. "
     "(2) OpenStreetMap Overpass API if SQLite returns nothing. "
     "(3) Google Places as a last resort. "
     "Facility type is matched to triage level: EMERGENCY → Hospitals only, "
     "MODERATE → PHC/CHC, MILD → Sub-Centre/PHC.",
     "No LLM — rule-based"),

    ("5", "Follow-up & Adherence Agent",
     "Care planner",
     "Generates a structured follow-up plan: when to return, what symptoms to watch for, "
     "what home care is appropriate, and red-flag triggers for immediate return. "
     "Follow-up interval is set by triage level (EMERGENCY → post-care, "
     "URGENT → 6 hours, MODERATE → 2 days, MILD → 4 days).",
     "Rule-based template"),

    ("6", "Health Worker Support Agent",
     "ASHA/ANM briefer",
     "Produces a structured briefing note for the frontline health worker "
     "(ASHA / ANM / CHW) before their home visit. "
     "Includes patient token, triage level, chief complaint, symptoms, "
     "recommended action, relevant government health schemes (PM-JAY, JSSK, NTEP), "
     "and escalation instructions.",
     "Rule-based template"),

    ("7", "Emergency Escalation Agent",
     "Life-saver node",
     "Triggered immediately when an emergency keyword is detected — bypasses all "
     "other agents. Shows 112 (National Emergency) and 108 (Ambulance) immediately. "
     "Retrieves first-aid instructions from the emergency_protocols ChromaDB collection. "
     "Response guaranteed in under 2 seconds. "
     "Falls back to a hardcoded static first-aid dict if ChromaDB is unavailable.",
     "No LLM — hardcoded + RAG fast path"),

    ("8", "Audit, Safety & Compliance Agent",
     "Terminal gate",
     "The last node in every pipeline path — always runs, even on pipeline failure. "
     "Runs the regex safety filter on the assembled response (blocks diagnosis patterns, "
     "prescription patterns). Injects the mandatory disclaimer if missing. "
     "Writes a complete audit log entry to SQLite (SHA-256 hashes, triage level, "
     "agent name, latency, safety result). "
     "The audit log is accessible only via the FastAPI /api/v1/audit endpoint.",
     "No LLM — code-enforced"),
]

for num, name, role, description, requirement in agent_details:
    p = doc.add_paragraph()
    r1 = p.add_run(f"Agent {num}: {name}  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    r2 = p.add_run(f"[{role}]")
    r2.italic = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.add_run(description).font.size = Pt(11)
    p2.paragraph_format.space_after = Pt(2)

    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Inches(0.3)
    r3 = p3.add_run(f"Requirement: {requirement}")
    r3.italic = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p3.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ── P.5 High-Level Architecture ──────────────────────────────────────────────
add_heading(doc, "P.5  High-Level Architecture", 2)
add_para(doc,
    "RuralCare AI is a six-layer system. Patient input enters at the Presentation Layer "
    "and flows down through Translation, Orchestration, Agent, Knowledge, and Data layers "
    "before a response is returned.",
    italic=True
)
doc.add_paragraph()

arch_layers = [
    ("Layer 1 — Presentation",
     "Streamlit UI (port 8501) for demo / evaluation.\n"
     "FastAPI REST API (port 8000) for programmatic access.\n"
     "Both call the same run_pipeline() function internally."),
    ("Layer 2 — Translation Middleware",
     "langdetect auto-detects patient's language.\n"
     "Google Translate converts input to English for LLM processing.\n"
     "Final response back-translated to patient's language before display."),
    ("Layer 3 — Orchestration",
     "LangGraph StateGraph manages 8 agent nodes.\n"
     "Pre-LLM emergency keyword check routes directly to Emergency agent.\n"
     "Fail-safe _node wrapper around every agent catches exceptions silently."),
    ("Layer 4 — Agent Pipeline",
     "8 specialist agents execute in sequence, sharing PatientState.\n"
     "Agents 7 (Emergency) and 8 (Audit) always execute regardless of path.\n"
     "PHI anonymisation: patient_token (PT-xxxxxxxx) used in all LLM prompts."),
    ("Layer 5 — Knowledge & Data",
     "ChromaDB: 6 vector collections, cosine similarity, multi-collection sweep.\n"
     "SQLite: 87 Tamil Nadu hospitals + audit logs + sessions + follow-up reminders.\n"
     "Auto-seeded on first boot in cloud deployments (Render / HF Spaces)."),
    ("Layer 6 — Infrastructure",
     "Local: Docker + docker-compose (Streamlit + FastAPI).\n"
     "Cloud: Render Starter plan + 1 GB Persistent Disk (~$7.25/month).\n"
     "Secrets: ANTHROPIC_API_KEY stored in Render dashboard, never in code.\n"
     "Source: github.com/jbobbypaul-cmyk/PGD-AIML-Examples-"),
]
for layer_name, layer_desc in arch_layers:
    p = doc.add_paragraph()
    r1 = p.add_run(layer_name + "  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    p.paragraph_format.space_after = Pt(2)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.add_run(layer_desc).font.size = Pt(11)
    p2.paragraph_format.space_after = Pt(8)

doc.add_paragraph()
add_heading(doc, "P.5.1  LangGraph Pipeline Flow", 3)
flow_text = (
    "Patient types symptoms (text, any language)\n"
    "  |\n"
    "  v\n"
    "[ Translation Middleware ]\n"
    "  Language detection --> Google Translate --> English text\n"
    "  |\n"
    "  v\n"
    "[ LangGraph Orchestrator ]\n"
    "  |\n"
    "  +--[ Emergency keyword detected? ]--YES--> Agent 7: Emergency Escalation\n"
    "  |                                               |\n"
    "  |                                               v\n"
    "  |                                         Agent 8: Audit & Safety --> RESPONSE\n"
    "  |\n"
    "  +--[ Normal ]-->\n"
    "       Agent 1: Symptom Intake      (LLM: extract structured symptoms)\n"
    "          |\n"
    "       Agent 2: Medical Triage      (LLM: EMERGENCY / URGENT / MODERATE / MILD)\n"
    "          |\n"
    "          +--[ EMERGENCY? ]-----------> Agent 7: Emergency Escalation\n"
    "          |\n"
    "       Agent 3: Medical RAG         (ChromaDB 6-collection sweep + LLM grounding)\n"
    "          |\n"
    "       Agent 4: Facility Lookup     (SQLite --> OSM --> Google Places)\n"
    "          |\n"
    "       Agent 5: Follow-up Plan      (Rule-based: watch-for, home care, return triggers)\n"
    "          |\n"
    "       Agent 6: Health Worker Note  (ASHA/ANM briefing + health schemes)\n"
    "          |\n"
    "       Agent 8: Audit & Safety      (Regex filter + disclaimer + SHA-256 audit log)\n"
    "          |\n"
    "          v\n"
    "  Back-translation to patient language\n"
    "  |\n"
    "  v\n"
    "RESPONSE displayed in Streamlit UI / returned as FastAPI JSON"
)
add_code_block(doc, flow_text)
doc.add_paragraph()

add_heading(doc, "P.5.2  Emergency Path Detail (< 2 seconds)", 3)
emergency_flow = (
    "'My father is unconscious and cannot breathe'\n"
    "  |\n"
    "  v\n"
    "detect_emergency()  <-- rule-based keyword scan, < 100ms, NO LLM call\n"
    "  |\n"
    "  MATCH: 'unconscious' + 'cannot breathe'\n"
    "  |\n"
    "  v\n"
    "Emergency Escalation Agent:\n"
    "  1. Show: CALL 112 (National Emergency) / CALL 108 (Ambulance)\n"
    "  2. Retrieve top-3 first-aid chunks from emergency_protocols ChromaDB collection\n"
    "  3. Show nearest hospital from SQLite facility cache\n"
    "  4. No LLM call -- hardcoded safety guaranteed\n"
    "  |\n"
    "  v\n"
    "Audit Safety Agent:\n"
    "  Write audit log entry (emergency_flag=True, latency < 2s)\n"
    "  |\n"
    "  v\n"
    "RESPONSE in < 2 seconds"
)
add_code_block(doc, emergency_flow)

add_heading(doc, "P.5.3  5-Layer Safety Architecture", 3)
safety_rows = [
    ("Layer 1", "Pre-LLM keyword scan",        "detect_emergency()",            "< 100ms, no LLM, rule-based"),
    ("Layer 2", "Emergency escalation agent",  "Agent 7",                       "Hardcoded 112/108 + RAG first-aid only"),
    ("Layer 3", "LLM system prompt",           "Every agent prompt",            "Explicitly forbids diagnosis and prescription"),
    ("Layer 4", "Regex output filter",         "run_safety_filter()",           "Blocks diagnosis/prescription patterns post-LLM"),
    ("Layer 5", "Audit safety gate",           "Agent 8 (terminal, always runs)","Disclaimer injection + SHA-256 audit log"),
]
add_table(doc, ["Layer", "Name", "Implementation", "Guarantee"],
          safety_rows, col_widths=[0.5, 1.8, 2.0, 2.3])

doc.add_page_break()

add_heading(doc, "P.6  Deployment Status", 2)
deploy_status_rows = [
    ("GitHub", "Source code + all configs", "github.com/jbobbypaul-cmyk/PGD-AIML-Examples-", "LIVE"),
    ("Local Dev", "Docker-compose or Streamlit direct", "localhost:8501 (UI) + localhost:8000 (API)", "READY"),
    ("Render", "Starter plan + 1GB Persistent Disk",
     "render.yaml pre-configured; set ANTHROPIC_API_KEY in dashboard", "READY TO DEPLOY"),
    ("HF Spaces", "Free Streamlit Space",
     "README.md front matter configured; set ANTHROPIC_API_KEY as HF secret", "READY TO DEPLOY"),
    ("Kaggle", "In-notebook demo pipeline",
     "notebooks/ruralcare_ai_demo.ipynb; set LLM_PROVIDER=claude in secrets", "READY"),
]
add_table(doc,
    ["Platform", "Setup", "Details", "Status"],
    deploy_status_rows,
    col_widths=[1.2, 1.8, 3.0, 1.2])
doc.add_page_break()


# =============================================================================
# 1. EXECUTIVE SUMMARY
# =============================================================================
add_heading(doc, "1. Executive Summary", 1)
add_para(doc,
    "RuralCare AI is an open-source, multilingual AI agent system designed to "
    "bridge the healthcare access gap in rural and underserved communities. It acts as a "
    "first-mile digital health assistant — available 24/7, multilingual, and grounded in "
    "verified public health knowledge from WHO and India's National Health Mission (NHM)."
)
doc.add_paragraph()
add_callout(doc,
    "IMPORTANT: RuralCare AI is NOT a diagnostic tool and does NOT replace a licensed "
    "healthcare professional. All outputs are for informational guidance only.",
    bg="FDECEA"
)
doc.add_paragraph()

add_heading(doc, "What RuralCare AI Does", 2)
bullets = [
    "Collects and structures patient-reported symptoms via text input",
    "Classifies urgency: Emergency / Urgent / Moderate / Mild",
    "Retrieves evidence-based health guidance from WHO and NHM documents (RAG)",
    "Locates the nearest appropriate government clinic or hospital",
    "Generates a personalised follow-up care plan",
    "Produces a briefing note for ASHA / ANM / CHW health workers",
    "Auto-escalates life-threatening emergencies with first-aid instructions",
    "Logs every interaction with a full audit trail for safety and compliance",
]
for b in bullets:
    add_bullet(doc, b)
doc.add_page_break()


# =============================================================================
# 2. KEY HIGHLIGHTS
# =============================================================================
add_heading(doc, "2. Key Highlights", 1)

highlights = [
    ("Multilingual Support",
     "English, Hindi, Tamil, Bengali, Telugu, Kannada — patient input detected and "
     "translated automatically; response delivered in the patient's language."),
    ("8-Agent LangGraph Pipeline",
     "LangGraph state machine with 8 specialist agents: Symptom Intake, Medical Triage, "
     "RAG Knowledge, Facility Lookup, Follow-up, Health Worker Briefing, Emergency Escalation, "
     "and Audit & Safety."),
    ("RAG-Grounded Responses Only",
     "Every health answer is retrieved from 6 ChromaDB collections (WHO, NHM, emergency "
     "protocols, symptom mapping, drug info, health schemes). The LLM is forbidden from "
     "improvising medical facts."),
    ("Sub-100ms Emergency Detection",
     "Emergency red-flag detection runs synchronously before any LLM call — purely rule-based "
     "keyword scan. Life-threatening symptoms trigger immediate escalation with 112/108 contacts."),
    ("87 Tamil Nadu Government Hospitals",
     "Government hospitals, CHCs, and PHCs across all 38 Tamil Nadu districts pre-loaded in "
     "SQLite, with OpenStreetMap and Google Places as fallback layers."),
    ("Privacy by Design",
     "Patient names and contact details never enter LLM prompts. Anonymised tokens "
     "(PT-xxxxxxxx) used throughout. Audit logs store SHA-256 hashes only."),
    ("Live Document Upload",
     "Hospital CSV/Excel files load directly into the facility database; medical PDF/TXT files "
     "are chunked and indexed into ChromaDB — both instantly searchable."),
    ("Local LLM — No API Cost for Local Dev",
     "Default LLM is Ollama llama3.2 running locally — free, offline, no API key needed. "
     "Switchable to Claude, GPT-4o, or Gemini by changing one environment variable."),
    ("5-Layer Safety Architecture",
     "Pre-LLM keyword detection → Emergency agent → LLM system prompt constraints → "
     "Regex output filter → Audit safety agent. No diagnosis or prescription possible."),
    ("Cloud Deployment Ready",
     "render.yaml configured for Render (Starter plan ~$7.25/month). "
     "HF Spaces front matter in README.md. Auto-seeds ChromaDB and SQLite on first boot."),
]

for title_text, desc in highlights:
    p = doc.add_paragraph()
    r_title = p.add_run(title_text + "  ")
    r_title.bold = True
    r_title.font.size = Pt(12)
    r_title.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    r_desc = p.add_run(desc)
    r_desc.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()


# =============================================================================
# 3. HOW TO USE IT
# =============================================================================
add_heading(doc, "3. How to Use RuralCare AI", 1)

add_heading(doc, "3.1  Quick Setup (Local)", 2)
add_para(doc, "Prerequisites: Python 3.10+, Ollama installed (ollama.ai)")
doc.add_paragraph()
steps = [
    ("Step 1 — Pull LLM models",
     "ollama pull llama3.2\nollama pull nomic-embed-text"),
    ("Step 2 — Install dependencies",
     "pip install -r requirements.txt"),
    ("Step 3 — Configure environment",
     "copy .env.example .env\n# Edit .env: set SQLITE_PATH and CHROMA_DB_PATH to absolute paths"),
    ("Step 4 — Seed Tamil Nadu hospitals",
     "python scripts/seed_tn_hospitals.py"),
    ("Step 5 — Load sample knowledge base",
     "python -m app.rag.document_loader --sample"),
    ("Step 6 — Run Streamlit demo UI",
     "$env:PYTHONUTF8='1'\npython -m streamlit run app/streamlit_app.py"),
    ("Step 7 — (Optional) Run FastAPI",
     "$env:PYTHONUTF8='1'\npython -m uvicorn app.main:app --reload --port 8000"),
]
for step_label, step_code in steps:
    add_para(doc, step_label, bold=True)
    add_code_block(doc, step_code)
    doc.add_paragraph()

add_heading(doc, "3.2  Using the Streamlit UI (port 8501)", 2)
ui_rows = [
    ("Sidebar",       "Select language (English/Hindi/Tamil/Bengali/Telugu/Kannada) and Quick Demo scenario buttons"),
    ("Main input",    "Type symptoms in the text box; enter District and State for facility lookup; click 'Get Health Guidance'"),
    ("Results panel", "Triage Level metric, Emergency Alert, Health Guidance, Facility cards, Follow-up Plan, Health Worker Briefing"),
    ("Upload panel",  "Expand 'Upload Document' to add hospital CSV files or medical PDF/TXT knowledge documents"),
]
add_table(doc, ["Area", "What to do"], ui_rows, col_widths=[1.8, 4.8])
doc.add_paragraph()

add_heading(doc, "3.3  Quick Demo Scenarios", 2)
demo_rows = [
    ("Demo 1", "Mild",        "Headache + tiredness",           "MILD — rest and monitor at home"),
    ("Demo 2", "Moderate",    "Fever 3 days + body aches",      "MODERATE — visit PHC within 48 hours"),
    ("Demo 3", "Urgent",      "Chest pain + 2 weeks cough",     "URGENT — see a doctor within 2-4 hours"),
    ("Demo 4", "Emergency",   "Unconscious, cannot breathe",    "EMERGENCY — call 112 immediately"),
    ("Demo 5", "Multilingual","Hindi: fever + headache input",  "MODERATE — response returned in Hindi"),
]
add_table(doc, ["Button", "Triage", "Input", "Expected Output"], demo_rows,
          col_widths=[0.8, 1.0, 2.2, 2.6])
doc.add_paragraph()

add_heading(doc, "3.4  FastAPI Endpoints (port 8000)", 2)
add_para(doc, "Open http://localhost:8000/docs for interactive Swagger UI.")
api_rows = [
    ("GET",  "/health",                   "Liveness probe — confirms server is running"),
    ("POST", "/api/v1/intake",            "Submit text symptoms — runs full 8-agent pipeline"),
    ("GET",  "/api/v1/facilities",        "Facility lookup by district, state, triage level"),
    ("GET",  "/api/v1/audit/{session_id}","Retrieve audit log for a specific session"),
    ("GET",  "/api/v1/audit",             "List recent audit entries (paginated, last 20 default)"),
    ("POST", "/api/v1/followup",          "Schedule a follow-up reminder"),
]
add_table(doc, ["Method", "Endpoint", "Purpose"], api_rows, col_widths=[0.7, 2.8, 3.1])
doc.add_page_break()


# =============================================================================
# 4. ARCHITECTURE OVERVIEW
# =============================================================================
add_heading(doc, "4. Architecture Overview", 1)

add_para(doc,
    "RuralCare AI is a multi-agent system orchestrated by LangGraph. Patient input flows "
    "through 8 specialist agents in sequence. Each agent reads from and writes to a shared "
    "PatientState object. A fail-safe wrapper around every agent ensures that a single "
    "failure never crashes the pipeline."
)
doc.add_paragraph()

add_heading(doc, "4.1  System Layers", 2)
arch_rows = [
    ("Patient Interface",   "Streamlit UI (port 8501)  |  FastAPI REST API (port 8000)"),
    ("Translation",         "langdetect (language detection) + Google Translate (input/response)"),
    ("Orchestration",       "LangGraph StateGraph — 8 nodes, conditional emergency routing, fail-safe _node wrapper"),
    ("LLM Layer",           "Ollama llama3.2 (local default)  |  Claude / GPT-4o / Gemini (cloud, via LLM_PROVIDER env)"),
    ("RAG / Knowledge",     "ChromaDB PersistentClient — 6 collections, cosine similarity, nomic-embed-text / sentence-transformers"),
    ("Facility Lookup",     "SQLite NHM TN cache (87 hospitals) --> OpenStreetMap Overpass --> Google Places (3-level hybrid)"),
    ("Database",            "SQLite (demo) — sessions, audit_logs, facility_cache, followup_reminders"),
    ("Safety",              "5-layer: pre-LLM keyword scan --> emergency agent --> LLM prompt --> regex filter --> audit gate"),
]
add_table(doc, ["Layer", "Technology"], arch_rows, col_widths=[1.8, 4.8])
doc.add_paragraph()

add_heading(doc, "4.2  The 8 Agents", 2)
agents_rows = [
    ("1", "Symptom Intake",        "Extracts structured symptoms from free-form text. Runs emergency keyword check BEFORE any LLM call."),
    ("2", "Medical Triage",        "Classifies urgency: EMERGENCY / URGENT / MODERATE / MILD using WHO ICD-11 principles."),
    ("3", "Medical RAG",           "Queries 6 ChromaDB collections, assembles context, generates grounded health guidance via LLM."),
    ("4", "Appointment & Facility","3-level hybrid lookup: NHM TN SQLite --> OSM Overpass --> Google Places. Returns top 3 facilities."),
    ("5", "Follow-up & Adherence", "Generates structured follow-up plan: watch-for items, home care, return-immediately triggers."),
    ("6", "Health Worker Support", "Produces ASHA/ANM briefing note with patient token, triage, symptoms, recommended action, health schemes."),
    ("7", "Emergency Escalation",  "Triggered pre-LLM on keyword detection. Shows 112/108, first-aid from RAG, nearest hospital. < 2 seconds."),
    ("8", "Audit, Safety & Compliance","Final gate: regex safety filter, disclaimer injection, SHA-256 audit log. Always runs, even on failure."),
]
add_table(doc, ["#", "Agent", "Role"], agents_rows, col_widths=[0.3, 1.9, 4.4])
doc.add_paragraph()

add_heading(doc, "4.3  RAG Pipeline", 2)
rag_rows = [
    ("Embedding model",  "Ollama nomic-embed-text (~768-dim) for local  |  sentence-transformers for cloud (Render/HF)"),
    ("Vector database",  "ChromaDB PersistentClient — embedded, no server needed, cosine similarity"),
    ("Collections",      "6: who_health_guidelines, nhm_india_protocols, symptom_disease_mapping, drug_information_basic, emergency_protocols, regional_health_schemes"),
    ("Chunking",         "Custom chunk_text() — 512-char window, 64-char overlap, separator-aware (no spacy dependency)"),
    ("Retrieval",        "Multi-collection cosine sweep in priority order; deduplicated by content hash; top_k=5"),
    ("Score threshold",  "None — Ollama cosine scores are not normalised 0-1; volume cap (k=5) used instead"),
    ("Grounding rule",   "LLM uses ONLY retrieved context — parametric knowledge forbidden for medical content"),
    ("Auto-seed",        "On cloud first boot: seeds 6 collections with WHO/NHM demo documents + 87 TN hospitals"),
]
add_table(doc, ["Aspect", "Detail"], rag_rows, col_widths=[1.8, 4.8])
doc.add_page_break()


# =============================================================================
# 5. SAFETY & GUARDRAILS
# =============================================================================
add_heading(doc, "5. Safety & Guardrails", 1)

add_callout(doc,
    "These rules are non-negotiable and are enforced at code level — not just by prompt instructions.",
    bg="FDECEA"
)
doc.add_paragraph()

add_heading(doc, "5.1  What RuralCare AI Will Never Do", 2)
never_rows = [
    ("Diagnose",         "Will never say 'You have malaria' or any equivalent — blocked by regex filter"),
    ("Prescribe",        "Will never recommend prescription drugs or dosages — blocked by regex filter"),
    ("Skip disclaimer",  "Every patient response includes the safety disclaimer — injected by code, not LLM"),
    ("Delay emergency",  "Emergency contacts shown immediately — before any other output, before LLM call"),
    ("Store PHI",        "Patient names/contacts never stored; anonymised tokens (PT-xxxxxxxx) only"),
    ("Improvise medicine","LLM forbidden from using training knowledge for medical facts — RAG-only"),
]
add_table(doc, ["Prohibited Action", "Enforcement"], never_rows, col_widths=[2.0, 4.6])
doc.add_paragraph()

add_heading(doc, "5.2  7 AI Guardrail Principles", 2)
guardrail_rows = [
    ("Non-maleficence",   "LLM system prompt constraints + independent regex safety filter on every output"),
    ("Bounded scope",     "Disclaimer appended by code; safe fallback shown when no RAG context found"),
    ("Grounding",         "LLM restricted to retrieved chunks only; rag_sources tracked per response"),
    ("Proportionality",   "EMERGENCY bypasses LLM entirely — hardcoded 112/108 + first-aid shown immediately"),
    ("Privacy by design", "patient_token in all LLM calls; SHA-256 hashes in audit log; PHI scrubbed from outputs"),
    ("Auditability",      "audit_safety_agent is terminal node — always runs; every agent writes its log entry"),
    ("Fail safe",         "_node wrapper catches exceptions; errors produce safe fallback, never wrong response"),
]
add_table(doc, ["Principle", "How it is enforced"], guardrail_rows, col_widths=[1.8, 4.8])
doc.add_paragraph()

add_heading(doc, "5.3  Emergency Red-Flag Categories (40+ keywords)", 2)
emergency_rows = [
    ("Cardiovascular / Respiratory", "chest pain, cannot breathe, difficulty breathing, heart attack, shortness of breath"),
    ("Neurological",                 "unconscious, seizure, stroke, sudden vision loss, sudden severe headache, loss of consciousness"),
    ("Severe Bleeding",              "severe bleeding, blood vomiting, coughing blood, bleeding that won't stop"),
    ("Poisoning / Envenomation",     "snake bite, poisoning, swallowed poison, insecticide ingestion, rat poison"),
    ("Obstetric Emergency",          "heavy vaginal bleeding during pregnancy, eclampsia, baby not moving"),
    ("Paediatric Emergency",         "child not breathing, baby turning blue, child not responding"),
]
add_table(doc, ["Category", "Example Keywords"], emergency_rows, col_widths=[2.2, 4.4])
doc.add_page_break()


# =============================================================================
# 6. TECHNOLOGY STACK
# =============================================================================
add_heading(doc, "6. Technology Stack", 1)

tech_rows = [
    ("Frontend",            "Streamlit 1.35+",              "Demo UI on port 8501; multilingual text input, document upload"),
    ("Backend API",         "FastAPI + uvicorn",             "6 REST endpoints, auto Swagger docs at /docs, port 8000"),
    ("Agent Orchestration", "LangGraph",                     "8-node state graph, conditional routing, fail-safe _node wrappers"),
    ("RAG Framework",       "LangChain",                     "Retriever chains, prompt templates, LCEL composition"),
    ("Vector DB",           "ChromaDB PersistentClient",     "6 collections, cosine similarity, embedded (no server)"),
    ("Embeddings (local)",  "Ollama nomic-embed-text",       "Local ~768-dim; free, offline, no API key needed"),
    ("Embeddings (cloud)",  "sentence-transformers",         "all-MiniLM-L6-v2; used on Render/HF Spaces; EMBED_PROVIDER=huggingface"),
    ("LLM (local default)", "Ollama llama3.2",               "Local inference, no API key, ~2GB download"),
    ("LLM (cloud)",         "Claude claude-sonnet-4-6",      "Set LLM_PROVIDER=claude + ANTHROPIC_API_KEY in .env or Render secrets"),
    ("LLM (alternatives)",  "GPT-4o / Gemini 2.0 Flash",    "Set LLM_PROVIDER=openai|gemini + respective API key"),
    ("Translation",         "Google Translate + langdetect", "Input to English + back-translation of response to patient language"),
    ("Facility data",       "SQLite + OSM + Google Places",  "3-level hybrid; 87 TN hospitals pre-seeded"),
    ("Chunking",            "Custom chunk_text()",           "512-char window, 64-char overlap; no spacy dependency"),
    ("Monitoring",          "LangSmith (optional)",          "Full pipeline traces; set LANGSMITH_API_KEY in .env"),
    ("Python",              "3.11 (Docker/cloud) / 3.10+",  "3.14 tested on Windows with PYTHONUTF8=1"),
    ("Deployment",          "Docker + Render + HF Spaces",  "render.yaml + Dockerfile + README front matter all configured"),
]
add_table(doc, ["Layer", "Technology", "Notes"], tech_rows, col_widths=[1.7, 2.0, 2.9])
doc.add_paragraph()

add_heading(doc, "Switching the LLM — One Line", 2)
add_code_block(doc,
    "# In .env — no code changes required:\n"
    "LLM_PROVIDER=ollama          # local, free (default)\n"
    "LLM_PROVIDER=claude          # requires ANTHROPIC_API_KEY\n"
    "LLM_PROVIDER=openai          # requires OPENAI_API_KEY\n"
    "LLM_PROVIDER=gemini          # requires GOOGLE_API_KEY"
)
doc.add_page_break()


# =============================================================================
# 7. BENEFITS
# =============================================================================
add_heading(doc, "7. Benefits", 1)

add_heading(doc, "For Patients", 2)
for b in [
    "Available 24/7 without travel to a clinic",
    "Works in their native language (6 Indian languages)",
    "Identifies emergencies immediately — no waiting",
    "Provides actionable guidance grounded in WHO/NHM documents",
    "Locates the nearest government facility with contact details",
    "Generates a simple follow-up plan to monitor recovery",
]:
    add_bullet(doc, b)

doc.add_paragraph()
add_heading(doc, "For Health Workers (ASHA / ANM / CHW)", 2)
for b in [
    "Automated structured briefing note ready before the home visit",
    "Patient token, triage level, chief complaint, watch-for items all pre-populated",
    "Relevant government health schemes (PM-JAY, JSSK, NTEP) flagged automatically",
    "Reduces manual note-taking and cognitive load during triage",
]:
    add_bullet(doc, b)

doc.add_paragraph()
add_heading(doc, "For Health Program Administrators", 2)
for b in [
    "Full audit trail on every interaction — SHA-256 hashed, append-only",
    "Add new hospitals via CSV upload — no developer required",
    "Add new medical knowledge via PDF upload — immediately retrievable in RAG",
    "Configurable LLM — local Ollama for zero cost or Claude for higher quality",
    "Cloud-ready: GitHub + Render + HF Spaces deployment fully configured",
]:
    add_bullet(doc, b)

doc.add_paragraph()
add_heading(doc, "Technical Benefits", 2)
for b in [
    "No diagnosis or prescription possible — enforced at code level, not just prompt level",
    "Privacy by design — patient identifiers never reach the LLM",
    "Fail-safe architecture — a single agent failure produces a safe fallback, not a crash",
    "Offline-capable with Ollama — runs without internet for LLM inference",
    "Modular — each agent is independently testable and replaceable",
]:
    add_bullet(doc, b)
doc.add_page_break()


# =============================================================================
# 8. PROJECT SCOPE
# =============================================================================
add_heading(doc, "8. Project Scope", 1)

add_heading(doc, "In Scope (MVP Complete)", 2)
in_scope_rows = [
    ("Symptom collection",    "Structured intake via text input (6 languages)"),
    ("Triage classification", "4 levels: Emergency / Urgent / Moderate / Mild"),
    ("Health information",    "RAG-grounded from 6 ChromaDB collections (WHO + NHM)"),
    ("Facility discovery",    "NHM Tamil Nadu (87 hospitals) + OpenStreetMap + Google Places"),
    ("Follow-up support",     "Follow-up plan with watch-for alerts and home care tips"),
    ("Health worker tools",   "Structured briefing note for ASHA/ANM/CHW"),
    ("Emergency escalation",  "Auto-escalation with 112/108 and first-aid from RAG"),
    ("Audit logging",         "Full interaction audit trail — SHA-256 hashes only"),
    ("Multilingual",          "6 Indian languages + English (UI + Google Translate pipeline)"),
    ("Demo deployment",       "Streamlit UI + FastAPI + Docker + Render + HF Spaces"),
    ("Document upload",       "Hospital CSV --> SQLite; Medical PDF/TXT --> ChromaDB"),
    ("Cloud auto-seed",       "ChromaDB + SQLite auto-seeded on first boot in cloud deployments"),
]
add_table(doc, ["Feature", "Implementation"], in_scope_rows, col_widths=[2.2, 4.4])
doc.add_paragraph()

add_heading(doc, "Out of Scope", 2)
out_scope_rows = [
    ("Medical diagnosis",        "Requires licensed physician — legally and ethically out of bounds"),
    ("Prescription generation",  "Requires licensed prescriber — hard out of scope"),
    ("Real-time telemedicine",   "Requires video infrastructure — future phase"),
    ("EHR integration",          "Requires HL7/FHIR compliance layer — future phase"),
    ("Drug interaction checking","Requires pharmacy-grade database — future phase"),
]
add_table(doc, ["Feature", "Reason"], out_scope_rows, col_widths=[2.2, 4.4])
doc.add_page_break()


# =============================================================================
# 9. DEPLOYMENT & ROADMAP
# =============================================================================
add_heading(doc, "9. Deployment & Roadmap", 1)

add_heading(doc, "Current Deployment Targets", 2)
deploy_rows = [
    ("Local dev",       "docker-compose up  OR  streamlit run app/streamlit_app.py",
     "Complete"),
    ("Kaggle Demo",     "Kaggle Notebook — notebooks/ruralcare_ai_demo.ipynb, LLM_PROVIDER=claude",
     "Ready"),
    ("Render",          "Starter plan + 1GB Persistent Disk (~$7.25/mo) — render.yaml pre-configured",
     "Ready to deploy"),
    ("Hugging Face",    "Streamlit Space — README.md front matter configured, LLM_PROVIDER=claude",
     "Ready to deploy"),
]
add_table(doc, ["Environment", "Method", "Status"], deploy_rows, col_widths=[1.5, 4.0, 1.2])
doc.add_paragraph()

add_heading(doc, "Render Deployment — Quick Reference", 2)
render_rows = [
    ("Plan",           "Starter (512 MB RAM, no spin-down)  +  1 GB Persistent Disk"),
    ("Cost",           "~$7.25/month — suitable for small evaluation group"),
    ("Start command",  "streamlit run app/streamlit_app.py --server.port=$PORT --server.address=0.0.0.0 --server.headless=true"),
    ("Health check",   "/_stcore/health"),
    ("Disk mount",     "/app/data"),
    ("Secret to set",  "ANTHROPIC_API_KEY — set in Render dashboard, NOT in render.yaml"),
    ("Auto-seed",      "ChromaDB + SQLite seeded on first boot from startup() in streamlit_app.py"),
]
add_table(doc, ["Setting", "Value"], render_rows, col_widths=[1.8, 4.8])
doc.add_paragraph()

add_heading(doc, "Implementation Roadmap Status", 2)
roadmap_rows = [
    ("Phase 0", "Foundation",            "Config, SQLite schema, safety filter, data models",          "COMPLETE"),
    ("Phase 1", "Core Agents",           "8 agents, LangGraph orchestrator, LLM factory",              "COMPLETE"),
    ("Phase 2", "RAG Pipeline",          "ChromaDB, chunk_text(), 6 collections, retriever",           "COMPLETE"),
    ("Phase 3", "Complete Pipeline",     "Facility agent, follow-up, health worker, 87 TN hospitals",  "COMPLETE"),
    ("Phase 4", "Frontend + Demo",       "Streamlit UI, document upload, FastAPI, demo buttons",       "COMPLETE"),
    ("Phase 5", "Multilingual",          "Translation + 6 languages; auto language detection",         "COMPLETE"),
    ("Phase 6", "Production Hardening",  "FastAPI running; JWT/rate-limiting/PostgreSQL pending",      "PARTIAL"),
    ("Phase 7", "Cloud Deployment",      "GitHub uploaded; render.yaml + HF Spaces config ready",      "READY"),
]
add_table(doc, ["Phase", "Name", "Scope", "Status"], roadmap_rows,
          col_widths=[0.6, 1.6, 3.2, 1.0])
doc.add_page_break()


# =============================================================================
# 10. TARGET USERS
# =============================================================================
add_heading(doc, "10. Target Users", 1)

add_heading(doc, "Primary Users", 2)
primary_rows = [
    ("Rural patient",     "18-70 years, limited English, basic smartphone",
     "Type symptoms in their language; receive triage guidance and facility details"),
    ("Caregiver",         "Family member managing a patient's care",
     "Use on behalf of patient; share follow-up plan and facility contact"),
    ("ASHA / ANM / CHW",  "Trained frontline health workers",
     "Receive structured briefing note before home visits; use health scheme info"),
]
add_table(doc, ["User", "Profile", "Primary Use"], primary_rows, col_widths=[1.5, 2.0, 3.1])
doc.add_paragraph()

add_heading(doc, "Secondary Users", 2)
secondary_rows = [
    ("District health officer", "Reviews aggregated triage data and emergency patterns"),
    ("Public health researcher", "Analyses de-identified symptom trends from audit log"),
    ("NGO health program staff", "Deploys RuralCare AI in field programs; uploads local hospital lists"),
    ("Evaluation leads",        "Assesses pipeline quality, safety guardrails, and architecture via Render demo"),
]
add_table(doc, ["User", "Use"], secondary_rows, col_widths=[2.2, 4.4])
doc.add_page_break()


# =============================================================================
# 11. SUCCESS METRICS
# =============================================================================
add_heading(doc, "11. Success Metrics", 1)

metrics_rows = [
    ("Triage accuracy",                "> 80%",  "> 92%"),
    ("Emergency escalation sensitivity","> 95%", "> 99%"),
    ("Average response latency",       "< 12s (Ollama local)", "< 3s (Claude API)"),
    ("Language coverage",              "6 languages",  "12+ languages"),
    ("Audit log completeness",         "100%",   "100%"),
    ("Safety filter false negatives",  "< 1%",   "< 0.1%"),
    ("System uptime (Render)",         "95%",    "99.9%"),
]
add_table(doc, ["Metric", "MVP Target", "Production Target"], metrics_rows,
          col_widths=[2.5, 2.0, 2.1])
doc.add_paragraph()


# =============================================================================
# 12. PROJECT DOCUMENTATION GUIDE
# =============================================================================
add_heading(doc, "12. Project Documentation Guide", 1)
add_para(doc, "The following documentation files are available in the project root (GitHub):")
doc_rows = [
    ("README.md",                "Quick start, tech stack, API endpoints, guardrail summary"),
    ("ARCHITECTURE.md",          "Complete system design — agents, RAG, chunking, FastAPI, privacy"),
    ("AGENT_DESIGN.md",          "Per-agent specs: role, inputs, outputs, system prompt, boundaries"),
    ("WORKFLOW.md",              "End-to-end step-by-step flow with code examples"),
    ("SAFETY_GUARDRAILS.md",     "7 guardrail principles, 5-layer safety architecture, compliance"),
    ("DATA_MODEL.md",            "All SQLite tables, Pydantic models, ChromaDB schema"),
    ("RAG_KNOWLEDGE_BASE.md",    "6 collections, chunking, retrieval strategy, user upload path"),
    ("TECH_STACK.md",            "Every technology choice with rationale and version notes"),
    ("PROJECT_SCOPE.md",         "In/out of scope, target users, constraints, acceptance criteria"),
    ("IMPLEMENTATION_ROADMAP.md","Phase 0-7 with completion status and quick-start commands"),
    ("KAGGLE_DEMO_PLAN.md",      "Notebook structure, 5 demo cases, Kaggle setup and constraints"),
    ("render.yaml",              "Render deployment config — plan, disk, env vars, start command"),
]
add_table(doc, ["File", "Contents"], doc_rows, col_widths=[2.4, 4.2])
doc.add_paragraph()


# =============================================================================
# DISCLAIMER PAGE
# =============================================================================
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run(
    "DISCLAIMER\n\n"
    "RuralCare AI is a first-level health information and triage support tool. "
    "It does NOT diagnose, prescribe, or replace a licensed healthcare professional. "
    "All outputs are for informational guidance only. "
    "Always consult a qualified doctor for medical decisions. "
    "In an emergency, call 112 (National Emergency) or 108 (Ambulance) immediately."
)
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)


# =============================================================================
# SAVE
# =============================================================================
output_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "RuralCare_AI_Documentation.docx"
)
doc.save(output_path)
print(f"Document saved: {output_path}")
